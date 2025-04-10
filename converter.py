import os
import sys
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from datetime import datetime
import docx
import PyPDF2
import re


class FileConverterGUI:
    def __init__(self, root):
        self.root = root
        self.root.title("Document Converter")
        self.root.geometry("600x400")

        self.input_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "input")
        self.output_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
        os.makedirs(self.input_dir, exist_ok=True)
        os.makedirs(self.output_dir, exist_ok=True)

        self.setup_gui()

    def setup_gui(self):
        style = ttk.Style()
        style.configure("TButton", padding=5)
        style.configure("TLabel", padding=5)

        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        title_label = ttk.Label(main_frame, text="Document Converter", font=("Helvetica", 16))
        title_label.grid(row=0, column=0, columnspan=2, pady=10)

        upload_btn = ttk.Button(main_frame, text="Upload Files", command=self.upload_files)
        upload_btn.grid(row=1, column=0, pady=5, padx=5)

        convert_btn = ttk.Button(main_frame, text="Convert Files", command=self.convert_files)
        convert_btn.grid(row=1, column=1, pady=5, padx=5)

        self.files_listbox = tk.Listbox(main_frame, width=70, height=15)
        self.files_listbox.grid(row=2, column=0, columnspan=2, pady=10)

        self.status_label = ttk.Label(main_frame, text="Ready")
        self.status_label.grid(row=3, column=0, columnspan=2)

        scrollbar = ttk.Scrollbar(main_frame, orient="vertical", command=self.files_listbox.yview)
        scrollbar.grid(row=2, column=2, sticky="ns")
        self.files_listbox.configure(yscrollcommand=scrollbar.set)

        dir_frame = ttk.Frame(main_frame)
        dir_frame.grid(row=4, column=0, columnspan=2, pady=10)

        ttk.Label(dir_frame, text=f"Input Directory: {self.input_dir}").pack()
        ttk.Label(dir_frame, text=f"Output Directory: {self.output_dir}").pack()

    def upload_files(self):
        files = filedialog.askopenfilenames(
            title="Select files to convert",
            filetypes=[("Document files", "*.pdf;*.doc;*.docx")]
        )

        if not files:
            return

        for file_path in files:
            filename = os.path.basename(file_path)
            destination = os.path.join(self.input_dir, filename)

            try:
                with open(file_path, 'rb') as src, open(destination, 'wb') as dst:
                    dst.write(src.read())
                self.files_listbox.insert(tk.END, filename)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to copy {filename}: {str(e)}")

        self.status_label.config(text=f"Uploaded {len(files)} files")

    def is_sentence_end(self, text):
        """Check if text ends with a sentence-ending punctuation."""
        return bool(re.search(r'[.!?][\'")\]]?\s*$', text))

    def is_continuation(self, current_line, next_line):
        """Check if the next line is a continuation of the current line."""
        if not next_line:
            return False
        if next_line[0].isupper() and not self.is_sentence_end(current_line):
            return True
        return not next_line[0].isupper() or current_line.endswith((',', ';', ':'))

    def process_text(self, text):
        """Process text with improved paragraph and structure detection."""
        lines = text.split('\n')
        processed_lines = []
        current_paragraph = []

        i = 0
        while i < len(lines):
            line = lines[i].strip()

            # Skip empty lines
            if not line:
                if current_paragraph:
                    processed_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
                processed_lines.append('')
                i += 1
                continue

            # Handle special markers (Issue, Brief, etc.)
            if re.match(r'^(Issue|Brief|Press Byte|Actionable)', line):
                if current_paragraph:
                    processed_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
                processed_lines.append(line)
                i += 1
                continue

            # Handle bullet points and numbered lists
            if line.startswith(('*', '-', '•', '●', '■')) or re.match(r'^\d+\.', line):
                if current_paragraph:
                    processed_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
                processed_lines.append(line)
                i += 1
                continue

            # Handle normal text
            current_paragraph.append(line)

            # Check if we should end the paragraph
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if not self.is_continuation(line, next_line):
                    if current_paragraph:
                        processed_lines.append(' '.join(current_paragraph))
                        current_paragraph = []

            i += 1

        # Handle any remaining text
        if current_paragraph:
            processed_lines.append(' '.join(current_paragraph))

        return '\n'.join(line for line in processed_lines if line)

    def clean_text(self, text):
        """Clean and normalize text while preserving structure."""
        # Normalize special characters
        text = text.replace('•', '*')
        text = text.replace('●', '*')
        text = text.replace('■', '*')
        text = text.replace('\u2013', '-')  # en dash
        text = text.replace('\u2014', '-')  # em dash
        text = text.replace('\u2018', "'")  # single quotes
        text = text.replace('\u2019', "'")
        text = text.replace('\u201c', '"')  # double quotes
        text = text.replace('\u201d', '"')

        # Remove extra whitespace while preserving structure
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)

        return text.strip()

    def extract_text_from_pdf(self, file_path):
        """Extract and format text from PDF files."""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = []

                for page in pdf_reader.pages:
                    text = page.extract_text()
                    if text.strip():
                        # Process each page's text
                        processed_text = self.process_text(text)
                        text_parts.append(processed_text)

                # Join pages and clean the final text
                full_text = '\n\n'.join(text_parts)
                return self.clean_text(full_text)
        except Exception as e:
            return f"Error processing PDF file: {str(e)}"

    def extract_text_from_doc(self, file_path):
        """Extract and format text from DOCX files."""
        try:
            doc = docx.Document(file_path)
            text_parts = []

            # Process paragraphs
            current_section = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    current_section.append(text)
                elif current_section:
                    text_parts.append(self.process_text('\n'.join(current_section)))
                    current_section = []

            if current_section:
                text_parts.append(self.process_text('\n'.join(current_section)))

            # Process tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text.append(' | '.join(row_text))
                if table_text:
                    text_parts.append('\n'.join(table_text))

            # Join all parts and clean the text
            full_text = '\n\n'.join(text_parts)
            return self.clean_text(full_text)
        except Exception as e:
            return f"Error processing DOCX file: {str(e)}"

    def convert_files(self):
        files = list(self.files_listbox.get(0, tk.END))
        if not files:
            messagebox.showwarning("Warning", "No files to convert!")
            return

        successful = 0
        failed = 0

        for filename in files:
            input_path = os.path.join(self.input_dir, filename)
            try:
                result = self.convert_to_txt(input_path)
                if "Successfully" in result:
                    successful += 1
                else:
                    failed += 1
                    messagebox.showerror("Error", f"Failed to convert {filename}: {result}")
            except Exception as e:
                failed += 1
                messagebox.showerror("Error", f"Error converting {filename}: {str(e)}")

        self.status_label.config(text=f"Converted {successful} files, {failed} failed")
        self.files_listbox.delete(0, tk.END)

        if successful > 0:
            messagebox.showinfo("Success",
                                f"Successfully converted {successful} files.\nOutput files are in: {self.output_dir}")

    def convert_to_txt(self, input_file):
        """Convert PDF or Word file to TXT."""
        if not os.path.exists(input_file):
            return f"Error: File {input_file} not found"

        file_lower = input_file.lower()
        filename = os.path.basename(input_file)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{os.path.splitext(filename)[0]}_{timestamp}.txt"
        output_file = os.path.join(self.output_dir, output_filename)

        try:
            if file_lower.endswith('.pdf'):
                text = self.extract_text_from_pdf(input_file)
            elif file_lower.endswith('.docx'):
                text = self.extract_text_from_doc(input_file)
            elif file_lower.endswith('.doc'):
                return f"Error: Old .doc format is not supported. Please save as .docx"
            else:
                return f"Error: Unsupported file format for {input_file}"

            if not text:
                return f"Error: No text could be extracted from {filename}"

            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(text)
            return f"Successfully converted {filename}"
        except Exception as e:
            return f"Error converting {filename}: {str(e)}"


def main():
    root = tk.Tk()
    app = FileConverterGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
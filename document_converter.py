import os
import re
from datetime import datetime
import docx
import PyPDF2

class DocumentConverter:
    def __init__(self, output_dir="output"):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)

    def is_sentence_end(self, text):
        return bool(re.search(r'[.!?][\'")\]]?\s*$', text))

    def is_continuation(self, current_line, next_line):
        if not next_line:
            return False
        if next_line[0].isupper() and not self.is_sentence_end(current_line):
            return True
        return not next_line[0].isupper() or current_line.endswith((',', ';', ':'))

    def process_text(self, text):
        lines = text.split('\n')
        processed_lines = []
        current_paragraph = []
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            if not line:
                if current_paragraph:
                    processed_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
                processed_lines.append('')
                i += 1
                continue
            if re.match(r'^(Issue|Brief|Press Byte|Actionable)', line):
                if current_paragraph:
                    processed_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
                processed_lines.append(line)
                i += 1
                continue
            if line.startswith(('*', '-', '•', '●', '■')) or re.match(r'^\d+\.', line):
                if current_paragraph:
                    processed_lines.append(' '.join(current_paragraph))
                    current_paragraph = []
                processed_lines.append(line)
                i += 1
                continue
            current_paragraph.append(line)
            if i + 1 < len(lines):
                next_line = lines[i + 1].strip()
                if not self.is_continuation(line, next_line):
                    if current_paragraph:
                        processed_lines.append(' '.join(current_paragraph))
                        current_paragraph = []
            i += 1
        if current_paragraph:
            processed_lines.append(' '.join(current_paragraph))
        return '\n'.join(line for line in processed_lines if line)

    def clean_text(self, text):
        text = text.replace('•', '*').replace('●', '*').replace('■', '*')
        text = text.replace('\u2013', '-').replace('\u2014', '-')
        text = text.replace('\u2018', "'").replace('\u2019', "'")
        text = text.replace('\u201c', '"').replace('\u201d', '"')
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        return text.strip()

    def extract_text_from_pdf(self, file_path):
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_parts = [self.process_text(page.extract_text()) for page in pdf_reader.pages if page.extract_text()]
            return self.clean_text('\n\n'.join(text_parts))
        except Exception as e:
            return f"Error processing PDF file: {str(e)}"

    def extract_text_from_doc(self, file_path):
        try:
            doc = docx.Document(file_path)
            text_parts, current_section = [], []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    current_section.append(text)
                elif current_section:
                    text_parts.append(self.process_text('\n'.join(current_section)))
                    current_section = []
            if current_section:
                text_parts.append(self.process_text('\n'.join(current_section)))
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            return self.clean_text('\n\n'.join(text_parts))
        except Exception as e:
            return f"Error processing DOCX file: {str(e)}"

    def convert_to_txt(self, input_file):
        if not os.path.exists(input_file):
            return f"Error: File {input_file} not found"
        file_lower = input_file.lower()
        filename = os.path.basename(input_file)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_filename = f"{os.path.splitext(filename)[0]}_{timestamp}.txt"
        output_path = os.path.join(self.output_dir, output_filename)

        try:
            if file_lower.endswith('.pdf'):
                text = self.extract_text_from_pdf(input_file)
            elif file_lower.endswith('.docx'):
                text = self.extract_text_from_doc(input_file)
            elif file_lower.endswith('.doc'):
                return f"Error: .doc format not supported. Convert to .docx."
            else:
                return f"Error: Unsupported file format {input_file}"
            if not text:
                return f"Error: No text extracted from {filename}"
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)
            return output_filename
        except Exception as e:
            return f"Error converting {filename}: {str(e)}"
